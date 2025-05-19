import streamlit as st
import feedparser
import requests
import pandas as pd
from datetime import datetime
from io import BytesIO
from textblob import TextBlob
from collections import Counter
from docx import Document
from bs4 import BeautifulSoup

st.set_page_config(page_title="📰 أداة الأخبار العربية الذكية", layout="wide")
st.title("🗞️ أداة إدارة وتحليل الأخبار (نسخة محسّنة + دعم مواقع عراقية)")

# التصنيفات
category_keywords = {
    "سياسة": ["رئيس", "وزير", "انتخابات", "برلمان", "سياسة", "رئاسة"],
    "رياضة": ["كرة", "لاعب", "مباراة", "دوري", "هدف"],
    "اقتصاد": ["سوق", "اقتصاد", "استثمار", "بنك", "مال"],
    "تكنولوجيا": ["تقنية", "تطبيق", "هاتف", "ذكاء", "برمجة"]
}

# --- التحليل والتصنيف ---
def summarize(text, max_words=25):
    return " ".join(text.split()[:max_words]) + "..."

def analyze_sentiment(text):
    polarity = TextBlob(text).sentiment.polarity
    if polarity > 0.1:
        return "😃 إيجابي"
    elif polarity < -0.1:
        return "😠 سلبي"
    else:
        return "😐 محايد"

def detect_category(text):
    for category, words in category_keywords.items():
        if any(word in text for word in words):
            return category
    return "غير مصنّف"

# --- Scrapers ---
def fetch_hathalyoum_news(keywords, date_from, date_to, chosen_category):
    url = "https://hathalyoum.net/"
    res = requests.get(url)
    soup = BeautifulSoup(res.content, "html.parser")
    news_list = []

    for article in soup.select(".newstitle a"):
        title = article.text.strip()
        link = article["href"]
        summary = title
        published_dt = datetime.today()

        full_text = title
        if keywords and not any(k.lower() in full_text.lower() for k in keywords):
            continue
        auto_category = detect_category(full_text)
        if chosen_category != "الكل" and auto_category != chosen_category:
            continue

        news_list.append({
            "source": "هَذا اليوم",
            "title": title,
            "summary": summary,
            "link": link,
            "published": published_dt,
            "image": "",
            "sentiment": analyze_sentiment(summary),
            "category": auto_category
        })
    return news_list

def fetch_iraqtoday_news(keywords, date_from, date_to, chosen_category):
    url = "https://iraqtoday.com/"
    res = requests.get(url)
    soup = BeautifulSoup(res.content, "html.parser")
    news_list = []

    for article in soup.select(".block-title a"):
        title = article.text.strip()
        link = article["href"]
        summary = title
        published_dt = datetime.today()

        full_text = title
        if keywords and not any(k.lower() in full_text.lower() for k in keywords):
            continue
        auto_category = detect_category(full_text)
        if chosen_category != "الكل" and auto_category != chosen_category:
            continue

        news_list.append({
            "source": "Iraq Today",
            "title": title,
            "summary": summary,
            "link": link,
            "published": published_dt,
            "image": "",
            "sentiment": analyze_sentiment(summary),
            "category": auto_category
        })
    return news_list

def fetch_moi_news(keywords, date_from, date_to, chosen_category):
    url = "https://moi.gov.iq/"
    res = requests.get(url)
    soup = BeautifulSoup(res.content, "html.parser")
    news_list = []

    for article in soup.select(".more-news .title-news a"):
        title = article.text.strip()
        link = article["href"]
        if not link.startswith("http"):
            link = "https://moi.gov.iq" + link
        summary = title
        published_dt = datetime.today()

        full_text = title
        if keywords and not any(k.lower() in full_text.lower() for k in keywords):
            continue
        auto_category = detect_category(full_text)
        if chosen_category != "الكل" and auto_category != chosen_category:
            continue

        news_list.append({
            "source": "وزارة الداخلية",
            "title": title,
            "summary": summary,
            "link": link,
            "published": published_dt,
            "image": "",
            "sentiment": analyze_sentiment(summary),
            "category": auto_category
        })
    return news_list

def fetch_presidency_news(keywords, date_from, date_to, chosen_category):
    url = "https://presidency.iq/"
    res = requests.get(url)
    soup = BeautifulSoup(res.content, "html.parser")
    news_list = []

    for article in soup.select(".article-title a"):
        title = article.text.strip()
        link = article["href"]
        summary = title
        published_dt = datetime.today()

        full_text = title
        if keywords and not any(k.lower() in full_text.lower() for k in keywords):
            continue
        auto_category = detect_category(full_text)
        if chosen_category != "الكل" and auto_category != chosen_category:
            continue

        news_list.append({
            "source": "رئاسة الجمهورية",
            "title": title,
            "summary": summary,
            "link": link,
            "published": published_dt,
            "image": "",
            "sentiment": analyze_sentiment(summary),
            "category": auto_category
        })
    return news_list

def fetch_asharq_iraq_news(keywords, date_from, date_to, chosen_category):
    url = "https://asharq.com/tags/%D8%A7%D9%84%D8%B9%D8%B1%D8%A7%D9%82/"
    res = requests.get(url)
    soup = BeautifulSoup(res.content, "html.parser")
    news_list = []

    for article in soup.select("a.Card"):
        title = article.text.strip()
        link = "https://asharq.com" + article["href"]
        summary = title
        published_dt = datetime.today()

        full_text = title
        if keywords and not any(k.lower() in full_text.lower() for k in keywords):
            continue
        auto_category = detect_category(full_text)
        if chosen_category != "الكل" and auto_category != chosen_category:
            continue

        news_list.append({
            "source": "الشرق - العراق",
            "title": title,
            "summary": summary,
            "link": link,
            "published": published_dt,
            "image": "",
            "sentiment": analyze_sentiment(summary),
            "category": auto_category
        })
    return news_list

def fetch_rt_iraq_news(keywords, date_from, date_to, chosen_category):
    url = "https://arabic.rt.com/focuses/10744-العراق/"
    res = requests.get(url)
    soup = BeautifulSoup(res.content, "html.parser")
    news_list = []

    for article in soup.select(".news-line a"):
        title = article.text.strip()
        link = "https://arabic.rt.com" + article["href"]
        summary = title
        published_dt = datetime.today()

        full_text = title
        if keywords and not any(k.lower() in full_text.lower() for k in keywords):
            continue
        auto_category = detect_category(full_text)
        if chosen_category != "الكل" and auto_category != chosen_category:
            continue

        news_list.append({
            "source": "RT - العراق",
            "title": title,
            "summary": summary,
            "link": link,
            "published": published_dt,
            "image": "",
            "sentiment": analyze_sentiment(summary),
            "category": auto_category
        })
    return news_list

# --- مصادر الأخبار ---
rss_feeds = {
    "هَذا اليوم": {"type": "scrape", "function": fetch_hathalyoum_news},
    "Iraq Today": {"type": "scrape", "function": fetch_iraqtoday_news},
    "وزارة الداخلية": {"type": "scrape", "function": fetch_moi_news},
    "رئاسة الجمهورية": {"type": "scrape", "function": fetch_presidency_news},
    "الشرق - العراق": {"type": "scrape", "function": fetch_asharq_iraq_news},
    "RT - العراق": {"type": "scrape", "function": fetch_rt_iraq_news}
}

# --- واجهة التحكم ---
col1, col2 = st.columns([1, 2])
with col1:
    selected_source = st.selectbox("🌐 اختر مصدر الأخبار:", list(rss_feeds.keys()))
    keywords_input = st.text_input("🔍 كلمات مفتاحية (مفصولة بفواصل):", "")
    keywords = [kw.strip() for kw in keywords_input.split(",")] if keywords_input else []
    category_filter = st.selectbox("📁 اختر التصنيف:", ["الكل"] + list(category_keywords.keys()))
    date_from = st.date_input("📅 من تاريخ:", datetime.today())
    date_to = st.date_input("📅 إلى تاريخ:", datetime.today())
    run = st.button("📥 عرض الأخبار")

# --- تحليل ونتائج ---
def export_to_word(news_list):
    doc = Document()
    for news in news_list:
        doc.add_heading(news['title'], level=2)
        doc.add_paragraph(f"المصدر: {news['source']}  |  التصنيف: {news['category']}")
        doc.add_paragraph(f"📅 التاريخ: {news['published'].strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"📄 التلخيص: {summarize(news['summary'])}")
        doc.add_paragraph(f"🔗 الرابط: {news['link']}")
        doc.add_paragraph(f"تحليل معنوي: {news['sentiment']}")
        doc.add_paragraph("-----")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_excel(news_list):
    df = pd.DataFrame(news_list)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    buffer.seek(0)
    return buffer

with col2:
    if run:
        source_data = rss_feeds[selected_source]
        news = source_data["function"](keywords, date_from, date_to, category_filter)

        if not news:
            st.warning("❌ لا توجد أخبار بهذه الشروط.")
        else:
            st.success(f"✅ تم العثور على {len(news)} خبر.")
            for item in news:
                with st.container():
                    st.markdown("----")
                    cols = st.columns([1, 4])
                    with cols[0]:
                        if item["image"]:
                            st.image(item["image"], use_column_width=True)
                    with cols[1]:
                        st.markdown(f"### 📰 {item['title']}")
                        st.markdown(f"📅 التاريخ: {item['published'].strftime('%Y-%m-%d')}")
                        st.markdown(f"📁 التصنيف: {item['category']}")
                        st.markdown(f"📄 التلخيص: {summarize(item['summary'])}")
                        st.markdown(f"🎯 التحليل: {item['sentiment']}")
                        st.markdown(f"[🌐 اقرأ المزيد ↗]({item['link']})")

            st.download_button("📄 تحميل كـ Word", data=export_to_word(news), file_name="news.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.download_button("📊 تحميل كـ Excel", data=export_to_excel(news), file_name="news.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

            st.markdown("### 🔠 أكثر الكلمات تكرارًا:")
            all_text = " ".join([n['summary'] for n in news])
            words = [word for word in all_text.split() if len(word) > 3]
            word_freq = Counter(words).most_common(10)
            for word, freq in word_freq:
                st.markdown(f"- **{word}**: {freq} مرة")
